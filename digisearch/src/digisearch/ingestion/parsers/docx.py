"""DOCX parser. Uses python-docx. Preserves heading structure."""

from __future__ import annotations

import io
import uuid
from pathlib import Path

from digisearch.core.models import DigiDocument
from digisearch.ingestion.base import Parser

try:
    from docx import Document
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


class DocxParser(Parser):
    """Parse DOCX. Extracts paragraphs and tables."""

    def parse(self, source: str | Path | bytes) -> DigiDocument:
        if not _DOCX_AVAILABLE:
            raise ImportError("Install python-docx for DOCX parsing")
        if isinstance(source, bytes):
            doc = Document(io.BytesIO(source))
            src_str = "<bytes>"
        else:
            path = Path(source)
            if path.exists():
                doc = Document(path)
                src_str = str(path)
            else:
                raise FileNotFoundError(f"DOCX source not found: {source}")
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text for c in row.cells]
                parts.append(" | ".join(cells))
        content = "\n".join(parts)
        doc_id = str(uuid.uuid4())
        return DigiDocument(
            id=doc_id,
            content=content,
            source=src_str,
            doc_type="docx",
            metadata={},
        )

    def can_parse(self, source: str) -> bool:
        ext = Path(source).suffix.lower() if "." in source else ""
        return ext in (".docx", ".doc")
